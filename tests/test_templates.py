"""Templates: discovery, and building a document on top of one."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from multycapture import paths, templates
from multycapture.generate import generate_docx

from conftest import SESSION_ID


@pytest.fixture
def templates_folder(monkeypatch, tmp_path) -> Path:
    """An empty templates directory that discovery will look in."""
    folder = tmp_path / "templates"
    folder.mkdir()
    monkeypatch.setattr(paths, "templates_dir", lambda: folder)
    return folder


def _write_template(folder: Path, name: str, *paragraphs: str) -> Path:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    path = folder / f"{name}.docx"
    doc.save(str(path))
    return path


def _text_of(path: Path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs if p.text.strip()]


# --------------------------------------------------------------------------- #
# discovery
# --------------------------------------------------------------------------- #
def test_no_templates_is_empty_not_an_error(templates_folder):
    assert templates.available() == []


def test_missing_folder_is_empty_not_an_error(monkeypatch, tmp_path):
    monkeypatch.setattr(paths, "templates_dir", lambda: tmp_path / "nope")
    assert templates.available() == []


def test_finds_docx_sorted_by_name(templates_folder):
    _write_template(templates_folder, "Zebra", "z")
    _write_template(templates_folder, "alpha", "a")
    assert [templates.label(p) for p in templates.available()] == ["alpha", "Zebra"]


def test_skips_word_lock_files(templates_folder):
    _write_template(templates_folder, "Report", "r")
    # Word leaves this behind next to any document it has open.
    (templates_folder / "~$Report.docx").write_bytes(b"not a document")
    assert [templates.label(p) for p in templates.available()] == ["Report"]


def test_ignores_non_docx(templates_folder):
    _write_template(templates_folder, "Real", "r")
    (templates_folder / "notes.txt").write_text("hello")
    (templates_folder / "old.doc").write_bytes(b"\xd0\xcf")
    assert [templates.label(p) for p in templates.available()] == ["Real"]


# --------------------------------------------------------------------------- #
# generating on top of a template
# --------------------------------------------------------------------------- #
def test_template_content_is_kept_and_steps_follow(capture, templates_folder, tmp_path):
    _, session_dir = capture
    tpl = _write_template(
        templates_folder, "Cover",
        "ACME S.p.A. — Manuale operativo", "Riservato, uso interno.",
    )

    out = generate_docx(str(session_dir), str(tmp_path / "with.docx"), template=str(tpl))
    text = _text_of(out)

    # the template's own content survived, in order, at the top
    assert text[0] == "ACME S.p.A. — Manuale operativo"
    assert text[1] == "Riservato, uso interno."
    # and the generated part comes after it
    assert "Captured Procedure" in text
    assert text.index("Captured Procedure") > text.index("Riservato, uso interno.")
    assert any(t.startswith("Step 1") for t in text)


def test_without_template_none_of_that_appears(capture, tmp_path):
    _, session_dir = capture
    out = generate_docx(str(session_dir), str(tmp_path / "plain.docx"))
    text = _text_of(out)
    assert not any("ACME" in t for t in text)
    assert text[0] == "Captured Procedure"


def test_template_with_content_gets_a_page_break(capture, templates_folder, tmp_path):
    _, session_dir = capture
    tpl = _write_template(templates_folder, "Cover", "Copertina")
    out = generate_docx(str(session_dir), str(tmp_path / "brk.docx"), template=str(tpl))
    xml = Document(str(out)).element.xml
    assert 'w:type="page"' in xml


def test_empty_template_gets_no_leading_blank_page(capture, templates_folder, tmp_path):
    """An empty template must not open the document with a blank page."""
    _, session_dir = capture
    tpl = _write_template(templates_folder, "Empty")  # no paragraphs at all
    out = generate_docx(str(session_dir), str(tmp_path / "empty.docx"), template=str(tpl))
    text = _text_of(out)
    assert text[0] == "Captured Procedure"
    assert 'w:type="page"' not in Document(str(out)).element.xml


def test_template_styles_reach_the_document(capture, templates_folder, tmp_path):
    """A style defined only in the template must exist in the output."""
    _, session_dir = capture
    doc = Document()
    doc.styles.add_style("AcmeBody", 1)  # 1 == WD_STYLE_TYPE.PARAGRAPH
    doc.add_paragraph("Copertina")
    tpl = templates_folder / "Styled.docx"
    doc.save(str(tpl))

    out = generate_docx(str(session_dir), str(tmp_path / "styled.docx"), template=str(tpl))
    assert "AcmeBody" in [s.name for s in Document(str(out)).styles]
