"""Steps that say what they caused.

A two-step session: click a button, and the screenshot of the next step shows a
dialog that was not there before. The document should say so.
"""

from __future__ import annotations

import datetime
import os

import pytest
from docx import Document
from PIL import Image, ImageDraw, ImageFont

from multycapture import ocr
from multycapture.capture import SessionReader
from multycapture.capture.session_writer import SessionWriter
from multycapture.generate import generate_docx
from multycapture.generate.condense import condense
from multycapture.generate.outcomes import read_outcomes
from multycapture.model import (
    CaptureConfig, ClickDetail, Event, EventType, MouseAction, MouseButton,
    Point, Rect, Session, ShotScope, TypeDetail, WindowInfo,
)

if not ocr.available():
    if os.environ.get("MULTYCAPTURE_REQUIRE_OCR"):
        raise RuntimeError(
            "MULTYCAPTURE_REQUIRE_OCR is set but tesseract is not installed"
        )
    pytestmark = pytest.mark.skip(reason="tesseract is not installed")

SIZE = (700, 500)


def _font(size: int):
    for path in (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        # Windows has no DejaVu; without this these tests skip for want of a
        # font even when tesseract is installed.
        r"C:\Windows\Fonts\arial.ttf",
    ):
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    pytest.skip("no scalable font available")


def _screen(with_dialog: bool = False, typed: str = "") -> Image.Image:
    image = Image.new("RGB", SIZE, (240, 240, 242))
    draw = ImageDraw.Draw(image)
    draw.rectangle([420, 400, 580, 448], fill=(0, 120, 215))
    draw.text((500, 413), "Save", font=_font(18), fill=(255, 255, 255), anchor="ma")
    if typed:
        draw.rectangle([60, 100, 500, 140], fill=(255, 255, 255), outline=(160, 160, 160))
        draw.text((70, 110), typed, font=_font(18), fill=(20, 20, 20))
    if with_dialog:
        draw.rectangle([150, 180, 560, 330], fill=(255, 255, 255), outline=(90, 90, 90))
        draw.text((190, 240), "Save changes?", font=_font(22), fill=(20, 20, 20))
    return image


def _session(tmp_path, screens, events_spec):
    """Build a session from (image, event-factory) pairs."""
    root = tmp_path / "captures"
    session = Session(
        id="session_20260812_150000",
        created_at=datetime.datetime.now().isoformat(),
        os="test", app_version="test",
        capture_config=CaptureConfig(shot_scope=ShotScope.MONITOR),
    )
    writer = SessionWriter(session, root).open()
    for image, make in zip(screens, events_spec):
        seq = writer.next_seq()
        shot = writer.save_shot(image, seq)
        writer.append_event(make(seq, shot))
    session.event_count = len(screens)
    writer.close()
    return writer.dir


def _click_at(x: int, y: int):
    """A click at a given spot.

    Two clicks in the *same* spot are one double-click as far as condense() is
    concerned, which would leave a single step and nothing to compare — so
    every fixture below clicks somewhere different, as a real procedure does.
    """
    def make(seq, shot):
        return Event(
            seq=seq, t=float(seq), ts=datetime.datetime.now().isoformat(),
            type=EventType.CLICK, screenshot=shot,
            mouse=Point(x, y), monitor=0,
            window=WindowInfo("Orders", "app", 1, Rect(0, 0, *SIZE)),
            mouse_rel=None, detail=ClickDetail(MouseButton.LEFT, MouseAction.DOWN),
        )
    return make


_SAVE = _click_at(500, 424)      # the Save button
_ELSEWHERE = _click_at(200, 300)  # somewhere else entirely


def _typing(text):
    def make(seq, shot):
        return Event(
            seq=seq, t=float(seq), ts=datetime.datetime.now().isoformat(),
            type=EventType.TYPE, screenshot=shot,
            mouse=Point(100, 120), monitor=0,
            window=WindowInfo("Orders", "app", 1, Rect(0, 0, *SIZE)),
            mouse_rel=None, detail=TypeDetail(text=text),
        )
    return make


def _outcomes_of(session_dir):
    reader = SessionReader(str(session_dir))
    session = reader.load_session()
    steps = condense(session, reader.events())
    return read_outcomes(reader, session, steps), steps


# --------------------------------------------------------------------------- #
def test_a_dialog_appearing_is_reported(tmp_path):
    session_dir = _session(
        tmp_path, [_screen(), _screen(with_dialog=True)], [_SAVE, _ELSEWHERE]
    )
    outcomes, steps = _outcomes_of(session_dir)

    assert steps[0].index in outcomes
    assert "changes" in outcomes[steps[0].index].lower()


def test_the_last_step_has_no_outcome(tmp_path):
    """There is no later screenshot to compare it against."""
    session_dir = _session(
        tmp_path, [_screen(), _screen(with_dialog=True)], [_SAVE, _ELSEWHERE]
    )
    outcomes, steps = _outcomes_of(session_dir)
    assert steps[-1].index not in outcomes


def test_an_unchanged_screen_says_nothing(tmp_path):
    session_dir = _session(tmp_path, [_screen(), _screen()], [_SAVE, _ELSEWHERE])
    outcomes, _ = _outcomes_of(session_dir)
    assert outcomes == {}


def test_typing_is_not_echoed_back(tmp_path):
    """The changed region after typing holds exactly what was typed.

    Repeating it would turn 'Type "ACME"' into 'Type "ACME". "ACME" appears.'
    """
    session_dir = _session(
        tmp_path,
        [_screen(), _screen(typed="ACME Industries")],
        [_typing("ACME Industries"), _SAVE],
    )
    outcomes, steps = _outcomes_of(session_dir)
    assert steps[0].index not in outcomes


def test_a_replaced_view_is_called_that(tmp_path):
    blank = Image.new("RGB", SIZE, (20, 20, 60))
    session_dir = _session(tmp_path, [_screen(), blank], [_SAVE, _ELSEWHERE])
    outcomes, steps = _outcomes_of(session_dir)
    assert outcomes[steps[0].index] == "The view changes."


def test_differently_sized_screenshots_are_skipped(tmp_path):
    """A moved window makes a pixel comparison meaningless."""
    session_dir = _session(
        tmp_path, [_screen(), _screen().resize((400, 300))], [_SAVE, _ELSEWHERE]
    )
    outcomes, _ = _outcomes_of(session_dir)
    assert outcomes == {}


def test_a_missing_screenshot_is_skipped_not_fatal(tmp_path):
    session_dir = _session(
        tmp_path, [_screen(), _screen(with_dialog=True)], [_SAVE, _ELSEWHERE]
    )
    for shot in (session_dir / "shots").glob("*.png"):
        shot.unlink()
    outcomes, _ = _outcomes_of(session_dir)
    assert outcomes == {}


# --------------------------------------------------------------------------- #
def test_the_document_says_what_happened(tmp_path):
    session_dir = _session(
        tmp_path, [_screen(), _screen(with_dialog=True)], [_SAVE, _ELSEWHERE]
    )
    out = generate_docx(str(session_dir), str(tmp_path / "with.docx"))
    text = " ".join(p.text for p in Document(str(out)).paragraphs)

    assert "Click" in text                    # what was done
    assert "changes" in text.lower()          # and what it caused


def test_outcomes_can_be_turned_off(tmp_path):
    session_dir = _session(
        tmp_path, [_screen(), _screen(with_dialog=True)], [_SAVE, _ELSEWHERE]
    )
    out = generate_docx(
        str(session_dir), str(tmp_path / "off.docx"), describe_outcomes=False
    )
    text = " ".join(p.text for p in Document(str(out)).paragraphs)
    assert "appears" not in text.lower()
    assert "Click" in text                    # the step itself is untouched
