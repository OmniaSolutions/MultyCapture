"""The tray's template rule.

One rule, both routes to a document: a configured template is used without
asking; no configured template means the chooser opens.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from multycapture import paths, templates
from multycapture.gui import tray as traymod

from conftest import SESSION_ID, wait_for_doc


@pytest.fixture
def templates_folder(monkeypatch, tmp_path) -> Path:
    folder = tmp_path / "templates"
    folder.mkdir()
    monkeypatch.setattr(paths, "templates_dir", lambda: folder)
    return folder


def _template(folder: Path, name: str, text: str = "Copertina") -> Path:
    doc = Document()
    doc.add_paragraph(text)
    path = folder / f"{name}.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def asked(monkeypatch):
    """Record chooser invocations; answer with whatever is queued."""
    calls: list[list[Path]] = []
    answer: list[tuple[bool, object]] = [(True, None)]

    def fake_ask(found, parent=None):
        calls.append(list(found))
        return answer[0]

    monkeypatch.setattr(traymod, "ask_template", fake_ask)
    return calls, answer


# --------------------------------------------------------------------------- #
# the setting
# --------------------------------------------------------------------------- #
def test_defaults_to_asking(tray):
    assert tray.template_mode == traymod._ASK


def test_setting_a_template_is_remembered(tray, templates_folder):
    from PySide6.QtCore import QSettings

    tpl = _template(templates_folder, "Aziendale")
    tray.set_template(traymod._FILE, str(tpl))
    tray.settings.sync()

    stored = QSettings("MultyCapture", "MultyCapture")
    assert stored.value("template_mode") == traymod._FILE
    assert stored.value("template_file") == str(tpl)


def test_menu_lists_templates_and_the_two_modes(tray, templates_folder):
    _template(templates_folder, "Aziendale")
    _template(templates_folder, "Breve")
    tray._rebuild_template_menu()
    labels = [a.text() for a in tray.template_menu.actions() if a.text()]
    assert "Ask every time" in labels
    assert "Blank document" in labels
    assert "Aziendale" in labels and "Breve" in labels


def test_menu_picks_up_a_template_added_while_running(tray, templates_folder):
    tray._rebuild_template_menu()
    assert "Tardivo" not in [a.text() for a in tray.template_menu.actions()]
    _template(templates_folder, "Tardivo")
    tray._rebuild_template_menu()  # what aboutToShow triggers
    assert "Tardivo" in [a.text() for a in tray.template_menu.actions()]


# --------------------------------------------------------------------------- #
# resolving — the rule itself
# --------------------------------------------------------------------------- #
def test_configured_template_is_used_without_asking(tray, templates_folder, asked):
    calls, _ = asked
    tpl = _template(templates_folder, "Aziendale")
    tray.set_template(traymod._FILE, str(tpl))

    proceed, chosen = tray._resolve_template()
    assert (proceed, chosen) == (True, str(tpl))
    assert calls == []  # nothing was asked


def test_blank_mode_is_used_without_asking(tray, templates_folder, asked):
    calls, _ = asked
    _template(templates_folder, "Aziendale")
    tray.set_template(traymod._BLANK)

    assert tray._resolve_template() == (True, None)
    assert calls == []


def test_unset_asks(tray, templates_folder, asked):
    calls, answer = asked
    tpl = _template(templates_folder, "Aziendale")
    answer[0] = (True, str(tpl))

    assert tray._resolve_template() == (True, str(tpl))
    assert len(calls) == 1
    assert calls[0] == [tpl]


def test_asks_even_with_a_single_template(tray, templates_folder, asked):
    """One template still gets a choice — a blank document stays an option."""
    calls, _ = asked
    _template(templates_folder, "Solo")
    tray._resolve_template()
    assert len(calls) == 1


def test_cancelling_the_chooser_stops_everything(tray, templates_folder, asked):
    _, answer = asked
    _template(templates_folder, "Aziendale")
    answer[0] = (False, None)
    assert tray._resolve_template() == (False, None)


def test_deleted_template_falls_back_to_asking(tray, templates_folder, asked):
    """A template configured then deleted must not break document generation."""
    calls, answer = asked
    tpl = _template(templates_folder, "Sparito")
    tray.set_template(traymod._FILE, str(tpl))
    tpl.unlink()
    answer[0] = (True, None)

    proceed, chosen = tray._resolve_template()
    assert (proceed, chosen) == (True, None)
    assert len(calls) == 1  # it asked instead of failing


# --------------------------------------------------------------------------- #
# both routes honour it
# --------------------------------------------------------------------------- #
def test_automatic_route_uses_the_configured_template(
    tray, capture, templates_folder, asked, documents_dir
):
    calls, _ = asked
    _, session_dir = capture
    tpl = _template(templates_folder, "Aziendale", "ACME — Manuale")
    tray.set_template(traymod._FILE, str(tpl))

    tray._maybe_auto_doc(str(session_dir), 2)
    assert wait_for_doc(tray)

    written = documents_dir / f"{SESSION_ID}.docx"
    text = [p.text for p in Document(str(written)).paragraphs if p.text.strip()]
    assert text[0] == "ACME — Manuale"
    assert calls == []  # the automatic route stayed automatic


def test_automatic_route_cancelled_generates_nothing(
    tray, capture, templates_folder, asked
):
    _, answer = asked
    _, session_dir = capture
    _template(templates_folder, "Aziendale")
    answer[0] = (False, None)

    tray._maybe_auto_doc(str(session_dir), 2)
    assert tray._doc_thread is None


def test_on_demand_route_asks_then_asks_for_the_folder(
    tray, capture, templates_folder, asked, tmp_path, monkeypatch
):
    from PySide6.QtWidgets import QFileDialog

    calls, answer = asked
    _, session_dir = capture
    tpl = _template(templates_folder, "Aziendale", "ACME — Manuale")
    answer[0] = (True, str(tpl))

    dest = tmp_path / "scelta"
    dest.mkdir()
    order: list[str] = []

    def fake_folder(parent, caption, directory, *a, **k):
        order.append("folder")
        return str(dest)

    monkeypatch.setattr(QFileDialog, "getExistingDirectory", fake_folder)

    def record_template(found, parent=None):
        order.append("template")
        return True, str(tpl)

    monkeypatch.setattr(traymod, "ask_template", record_template)

    tray.generate_doc()
    assert order == ["template", "folder"]  # what, then where
    assert wait_for_doc(tray)

    written = dest / f"{SESSION_ID}.docx"
    text = [p.text for p in Document(str(written)).paragraphs if p.text.strip()]
    assert text[0] == "ACME — Manuale"
