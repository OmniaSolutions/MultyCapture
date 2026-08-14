"""Instructions that name what was clicked.

Without OCR a step reads "Click in “Some App”." — true, and useless to anyone
following along. These tests build a session whose screenshot contains a real
button, click it, and check the label reaches the document.
"""

from __future__ import annotations

import datetime
import os

import pytest
from docx import Document
from PIL import Image, ImageDraw, ImageFont

from multycapture import ocr
from multycapture.capture import SessionReader
from multycapture.capture.session_writer import SessionWriter
from multycapture.generate import generate_docx
from multycapture.generate.condense import condense, raw_steps
from multycapture.generate.labels import read_labels, quoted
from multycapture.model import (
    CaptureConfig, ClickDetail, Event, EventType, MouseAction, MouseButton,
    Point, Rect, Session, ShotScope, WindowInfo,
)

if not ocr.available():
    if os.environ.get("MULTYCAPTURE_REQUIRE_OCR"):
        raise RuntimeError(
            "MULTYCAPTURE_REQUIRE_OCR is set but tesseract is not installed"
        )
    pytestmark = pytest.mark.skip(reason="tesseract is not installed")

BUTTON_AT = (300, 150)


def _font(size: int):
    for path in (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ):
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    pytest.skip("no scalable font available to render test text")


def _screen(label: str) -> Image.Image:
    """A window with one clearly-labelled primary button."""
    image = Image.new("RGB", (600, 300), (242, 242, 244))
    draw = ImageDraw.Draw(image)
    cx, cy = BUTTON_AT
    draw.rectangle([cx - 80, cy - 24, cx + 80, cy + 24], fill=(0, 120, 215))
    draw.text((cx, cy - 11), label, font=_font(20), fill=(255, 255, 255), anchor="ma")
    return image


@pytest.fixture
def clicked_session(tmp_path):
    """A one-click session whose screenshot shows a "Salva" button."""
    root = tmp_path / "captures"
    session = Session(
        id="session_20260812_120000",
        created_at=datetime.datetime.now().isoformat(),
        os="test",
        app_version="test",
        capture_config=CaptureConfig(shot_scope=ShotScope.MONITOR),
    )
    writer = SessionWriter(session, root).open()
    seq = writer.next_seq()
    shot = writer.save_shot(_screen("Salva"), seq)
    writer.append_event(Event(
        seq=seq, t=0.0, ts=datetime.datetime.now().isoformat(),
        type=EventType.CLICK, screenshot=shot,
        mouse=Point(*BUTTON_AT), monitor=0,
        window=WindowInfo("Gestione Ordini", "app", 1, Rect(0, 0, 600, 300)),
        mouse_rel=None,
        detail=ClickDetail(MouseButton.LEFT, MouseAction.DOWN),
    ))
    session.event_count = 1
    writer.close()
    return writer.dir


# --------------------------------------------------------------------------- #
def test_label_is_read_from_the_screenshot(clicked_session):
    reader = SessionReader(str(clicked_session))
    session = reader.load_session()
    events = reader.events()
    assert read_labels(reader, session, events) == {events[0].seq: "Salva"}


def test_condensed_instruction_names_the_button(clicked_session):
    reader = SessionReader(str(clicked_session))
    session = reader.load_session()
    events = reader.events()
    labels = read_labels(reader, session, events)

    with_label = condense(session, events, labels)[0].instruction
    without = condense(session, events)[0].instruction

    assert "Salva" in with_label
    assert with_label.startswith("Click “Salva”")
    assert "Salva" not in without  # the difference is the label, nothing else


def test_raw_steps_name_the_button_too(clicked_session):
    reader = SessionReader(str(clicked_session))
    session = reader.load_session()
    events = reader.events()
    labels = read_labels(reader, session, events)
    assert "Salva" in raw_steps(events, labels)[0].instruction


def test_the_document_says_what_was_clicked(clicked_session, tmp_path):
    out = generate_docx(str(clicked_session), str(tmp_path / "labelled.docx"))
    text = " ".join(p.text for p in Document(str(out)).paragraphs)
    assert "Click “Salva”" in text


def test_ocr_can_be_turned_off(clicked_session, tmp_path):
    out = generate_docx(
        str(clicked_session), str(tmp_path / "plain.docx"), read_click_labels=False
    )
    text = " ".join(p.text for p in Document(str(out)).paragraphs)
    assert "Salva" not in text
    assert "Click" in text  # still a usable instruction, just unnamed


def test_missing_screenshot_is_skipped_not_fatal(clicked_session):
    """A capture whose image is gone must still produce a document."""
    reader = SessionReader(str(clicked_session))
    session = reader.load_session()
    events = reader.events()
    for shot in (clicked_session / "shots").glob("*.png"):
        shot.unlink()
    assert read_labels(reader, session, events) == {}


def test_quoting_helper_is_empty_without_a_label():
    assert quoted(None) == ""
    assert quoted("") == ""
    assert quoted("Salva") == " “Salva”"
