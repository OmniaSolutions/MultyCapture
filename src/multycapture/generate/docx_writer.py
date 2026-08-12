"""Generate a Word (.docx) document from a captured session.

Each recorded event becomes a numbered step: an imperative instruction plus the
screenshot taken at that moment, with the click/scroll location highlighted.
"""

from __future__ import annotations

import datetime
import re
from pathlib import Path
from typing import Callable, Optional

from .. import paths
from ..capture import SessionReader
from ..model import Event, Session
from . import steps
from .condense import Step, condense, raw_steps
from .labels import read_labels
from .outcomes import read_outcomes
from .annotate import prepare_for_doc

# XML 1.0 forbids most C0 control chars (keep tab/newline/carriage-return).
# Captured window titles and typed text can contain these, so strip them.
_INVALID_XML = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f￾￿]")


def _safe(text: Optional[str]) -> str:
    return _INVALID_XML.sub("", text or "")


def _has_content(doc) -> bool:
    """Whether a template carries anything worth keeping on its own page."""
    if doc.tables or doc.inline_shapes:
        return True
    return any(p.text.strip() for p in doc.paragraphs)


def _fmt_when(iso: str) -> str:
    try:
        return datetime.datetime.fromisoformat(iso).strftime("%Y-%m-%d %H:%M")
    except (ValueError, TypeError):
        return iso or ""


def generate_docx(
    session_dir: str,
    out_path: Optional[str] = None,
    *,
    template: Optional[str] = None,
    title: Optional[str] = None,
    annotate: bool = True,
    condense_steps: bool = True,
    read_click_labels: bool = True,
    describe_outcomes: bool = True,
    rewrite: Optional[Callable[[list[Step], dict[int, str]], None]] = None,
    max_width: int = 1200,
    image_width_inches: float = 6.5,
) -> Path:
    """Build a .docx for ``session_dir``; return the written path.

    Without ``out_path`` the document is written to the user's documents
    directory (see :mod:`..paths`), named after the session.

    ``template`` is a .docx to build on. Its styles *and its content* carry
    over — a cover page or preamble in the template stays, and the generated
    steps are appended after it. Without one the document starts empty.

    By default the raw event stream is condensed into meaningful steps
    (:func:`condense`). Pass ``condense_steps=False`` for one step per event.

    ``rewrite`` is called with the finished steps and their labels before the
    document is written, to adjust the wording. It must not add, remove or
    reorder steps.

    ``describe_outcomes`` compares each step's screenshot with the next one
    and adds a sentence for what changed — "“Save changes?” appears." It needs
    the same tesseract as the labels and costs about 20 ms a step on top.

    ``read_click_labels`` reads the on-screen text each click landed on, so
    instructions name their target — "Click “Save”" rather than "Click".
    It costs roughly a tenth of a second per click and needs tesseract
    installed; without either it simply produces the unnamed wording.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image

    reader = SessionReader(session_dir)
    session = reader.load_session()
    events = reader.events()

    # Read what each click landed on, so instructions can name it. Done here,
    # at generation time, rather than while recording: reading an image per
    # event costs about a tenth of a second, which a recorder cannot spend
    # without dropping events.
    label_map = read_labels(reader, session, events) if read_click_labels else {}

    step_list = (
        condense(session, events, label_map) if condense_steps
        else raw_steps(events, label_map)
    )

    # What each step caused, from the screenshots either side of it. After
    # the steps exist, since it compares consecutive ones.
    outcomes = read_outcomes(reader, session, step_list) if describe_outcomes else {}

    # Optional pass over the finished steps — this is where an AI rewording
    # plugs in. It may change wording and nothing else; the caller is
    # responsible for enforcing that (see multycapture.ai.payload).
    if rewrite is not None:
        rewrite(step_list, label_map)

    if template:
        doc = Document(template)
        # Start the generated part on its own page, but only when the template
        # actually has something on the first one — otherwise an empty template
        # would open with a blank page.
        if _has_content(doc):
            doc.add_page_break()
    else:
        doc = Document()

    # ---- title + metadata -------------------------------------------------
    doc.add_heading(_safe(title or "Captured Procedure"), level=0)
    meta = doc.add_paragraph()
    detail = f"{len(step_list)} steps"
    if condense_steps:
        detail += f" (condensed from {len(events)} events)"
    meta_run = meta.add_run(
        _safe(f"{detail} · captured {_fmt_when(session.created_at)} · {session.os}")
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # ---- steps ------------------------------------------------------------
    for step in step_list:
        doc.add_heading(f"Step {step.index}", level=2)

        instr = doc.add_paragraph()
        instr.add_run(_safe(step.instruction)).bold = True

        outcome = outcomes.get(step.index)
        if outcome:
            said = doc.add_paragraph()
            run = said.add_run(_safe(outcome))
            run.italic = True
            run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

        _add_screenshot(
            doc, reader, session, step.event,
            annotate=annotate, max_width=max_width,
            width_inches=image_width_inches,
            Image=Image, Inches=Inches, Pt=Pt, RGBColor=RGBColor,
            WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
        )

    # ---- output path ------------------------------------------------------
    if out_path is None:
        # The user's Documents folder, not next to the capture: the capture
        # lives in application data, which nobody browses to.
        out = paths.documents_dir() / f"{session.id}.docx"
    else:
        out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _add_screenshot(
    doc, reader: SessionReader, session: Session, event: Event, *,
    annotate: bool, max_width: int, width_inches: float,
    Image, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
) -> None:
    path = reader.shot_path(event.screenshot)
    if path is None or not path.exists():
        note = doc.add_paragraph()
        run = note.add_run("[screenshot unavailable]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
        return

    point = None
    if annotate and steps.is_pointed(event):
        origin = steps.shot_origin(session, event)
        point = (event.mouse.x - origin.x, event.mouse.y - origin.y)

    with Image.open(path) as im:
        im.load()
        if point is not None:
            # clamp the highlight to the image so an off-window click still draws
            px = min(max(point[0], 0), im.width - 1)
            py = min(max(point[1], 0), im.height - 1)
            point = (px, py)
        buf = prepare_for_doc(im, point=point, max_width=max_width)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(buf, width=Inches(width_inches))

    # caption: app + relative time
    app = _safe(steps.app_label(event.window))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = caption.add_run(
        f"{app}  ·  t+{event.t:.1f}s" if app else f"t+{event.t:.1f}s"
    )
    crun.italic = True
    crun.font.size = Pt(8)
    crun.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
